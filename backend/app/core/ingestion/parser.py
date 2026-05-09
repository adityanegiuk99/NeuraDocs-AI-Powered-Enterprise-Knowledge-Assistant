"""
Document parser module.
Handles extraction of text and structured elements from PDF, DOCX, and TXT files.
Uses Unstructured for PDF parsing with layout analysis, python-docx for DOCX.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from app.utils.logging import get_logger

logger = get_logger(__name__)


@dataclass
class DocumentElement:
    """Represents a single structural element extracted from a document."""
    text: str
    element_type: str  # title, narrative, table, list_item, header, footer
    page_number: Optional[int] = None
    metadata: dict = field(default_factory=dict)


class DocumentParser:
    """
    Unified document parser that delegates to format-specific parsers.
    Extracts structured elements preserving document hierarchy.
    """

    def parse(self, file_path: str) -> list[DocumentElement]:
        """Parse a document and return structured elements."""
        path = Path(file_path)
        ext = path.suffix.lower()

        logger.info("parsing_document", file=path.name, type=ext)

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext == ".docx":
            return self._parse_docx(path)
        elif ext == ".txt":
            return self._parse_txt(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _parse_pdf(self, path: Path) -> list[DocumentElement]:
        """
        Parse PDF using PyPDF2 as the primary method.
        Falls back gracefully if unstructured is available for hi-res parsing.
        """
        elements = []

        try:
            # Try unstructured first for better layout analysis
            from unstructured.partition.pdf import partition_pdf

            raw_elements = partition_pdf(
                filename=str(path),
                strategy="fast",  # Use "hi_res" if OCR + layout needed
                include_page_breaks=True,
            )

            for elem in raw_elements:
                el_type = self._map_unstructured_type(type(elem).__name__)
                if elem.text and elem.text.strip():
                    elements.append(DocumentElement(
                        text=elem.text.strip(),
                        element_type=el_type,
                        page_number=getattr(elem.metadata, 'page_number', None),
                        metadata={
                            "source": "unstructured",
                            "category": type(elem).__name__,
                        },
                    ))

            logger.info("pdf_parsed_unstructured", elements=len(elements))

        except (ImportError, Exception) as e:
            # Fallback to PyPDF2
            logger.info("pdf_fallback_pypdf2", reason=str(e))
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    # Split into paragraphs
                    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                    for para in paragraphs:
                        el_type = self._detect_element_type(para)
                        elements.append(DocumentElement(
                            text=para,
                            element_type=el_type,
                            page_number=page_num,
                            metadata={"source": "pypdf2"},
                        ))

            logger.info("pdf_parsed_pypdf2", elements=len(elements))

        return elements

    def _parse_docx(self, path: Path) -> list[DocumentElement]:
        """Parse DOCX using python-docx, preserving headings and structure."""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        elements = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect element type from paragraph style
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                el_type = "title"
            elif "list" in style_name or "bullet" in style_name:
                el_type = "list_item"
            else:
                el_type = "narrative"

            elements.append(DocumentElement(
                text=text,
                element_type=el_type,
                metadata={"source": "python-docx", "style": para.style.name},
            ))

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._table_to_markdown(table)
            if table_text:
                elements.append(DocumentElement(
                    text=table_text,
                    element_type="table",
                    metadata={"source": "python-docx", "table_index": table_idx},
                ))

        logger.info("docx_parsed", elements=len(elements))
        return elements

    def _parse_txt(self, path: Path) -> list[DocumentElement]:
        """Parse plain text file, splitting into paragraphs."""
        text = path.read_text(encoding="utf-8", errors="replace")
        elements = []

        # Split on double newlines (paragraphs)
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

        for para in paragraphs:
            el_type = self._detect_element_type(para)
            elements.append(DocumentElement(
                text=para,
                element_type=el_type,
                metadata={"source": "plain_text"},
            ))

        logger.info("txt_parsed", elements=len(elements))
        return elements

    def _table_to_markdown(self, table) -> str:
        """Convert a docx table to markdown format for embedding."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) >= 2:
            # Insert markdown header separator after first row
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)

        return "\n".join(rows)

    def _map_unstructured_type(self, class_name: str) -> str:
        """Map Unstructured element class names to our simplified types."""
        mapping = {
            "Title": "title",
            "NarrativeText": "narrative",
            "Table": "table",
            "ListItem": "list_item",
            "Header": "header",
            "Footer": "footer",
            "FigureCaption": "narrative",
            "Address": "narrative",
        }
        return mapping.get(class_name, "narrative")

    def _detect_element_type(self, text: str) -> str:
        """Heuristic detection of element type from plain text."""
        stripped = text.strip()

        # Short, possibly title-like
        if len(stripped) < 100 and stripped.isupper():
            return "title"
        if len(stripped) < 100 and not stripped.endswith("."):
            return "title"
        # Bullet points
        if stripped.startswith(("- ", "• ", "* ", "1.", "2.", "3.")):
            return "list_item"
        # Table-like (has pipes or tabs)
        if "|" in stripped or "\t" in stripped:
            return "table"

        return "narrative"
